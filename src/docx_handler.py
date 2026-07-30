"""
docx_handler.py
Estampador de firmas en reportes de asistencia (DOCX / PDF).
Garantiza que la salida sea SIEMPRE un archivo PDF y que la firma encaje
exactamente dentro de las celdas sin alterar las dimensiones ni la estructura original.
"""

import os
import io
from PIL import Image as PILImage

try:
    import docx
    from docx.shared import Inches, Pt
except ImportError:
    docx = None

try:
    import pypdf
    from reportlab.pdfgen import canvas
except ImportError:
    pypdf = None

try:
    import fitz  # PyMuPDF
    import cv2
    import numpy as np
except ImportError:
    fitz = None
    cv2 = None
"""
docx_handler.py
Estampador de firmas en reportes de asistencia (DOCX / PDF).
Garantiza que la salida sea SIEMPRE un archivo PDF y que la firma encaje
exactamente dentro de las celdas sin alterar las dimensiones ni la estructura original.
"""

import os
import io
from PIL import Image as PILImage

try:
    import docx
    from docx.shared import Inches, Pt
except ImportError:
    docx = None

try:
    import pypdf
    from reportlab.pdfgen import canvas
except ImportError:
    pypdf = None

try:
    import fitz  # PyMuPDF
    import cv2
    import numpy as np
except ImportError:
    fitz = None
    cv2 = None
    np = None


class DocumentSigner:
    """Inserta la firma en archivos de asistencia y exporta SIEMPRE en formato PDF."""

    @classmethod
    def _detectar_casillas_firma_cv(cls, ruta_doc: str) -> list[tuple[float, float, float, float]]:
        """
        Analiza la primera página del PDF usando extracción de tabla de PyMuPDF y/o Visión por Computadora (OpenCV).
        Determina únicamente las celdas válidas pertenecientes a la columna 'FIRMA' en las filas de datos de asistencia,
        delimitando estrictamente la tabla para no colocar firmas en el pie de página ni sobre textos de resumen.
        Retorna lista de tuplas (x_pdf, y_pdf_bottom, width_pdf, height_pdf) en coordenadas PDF.
        """
        if not fitz:
            return []

        try:
            doc = fitz.open(ruta_doc)
            if len(doc) == 0:
                return []

            page = doc[0]
            rect = page.rect
            pdf_w, pdf_h = float(rect.width), float(rect.height)

            casillas_pdf = []

            # ── ENFOQUE 1: Extracción Estructurada de Tabla con PyMuPDF (Alta Precisión) ──
            tables = page.find_tables()
            if tables and len(tables.tables) > 0:
                for table in tables.tables:
                    extracted_matrix = table.extract()
                    if not extracted_matrix or len(extracted_matrix) < 2:
                        continue

                    # Identificar cuáles columnas contienen "FIRMA" en la cabecera o primeras filas
                    col_firma_indices = []
                    header_row_idx = 0

                    for r_idx in range(min(3, len(extracted_matrix))):
                        row_cells = [str(c or "").strip().upper() for c in extracted_matrix[r_idx]]
                        for c_idx, cell_text in enumerate(row_cells):
                            if "FIRMA" in cell_text and c_idx not in col_firma_indices:
                                col_firma_indices.append(c_idx)
                                header_row_idx = max(header_row_idx, r_idx)

                    if not col_firma_indices:
                        continue

                    # Recorrer las filas de la tabla posteriores a la cabecera
                    for r_idx in range(header_row_idx + 1, len(table.rows)):
                        row_cells_text = [str(c or "").strip().upper() for c in extracted_matrix[r_idx]]
                        full_row_str = " ".join(row_cells_text)

                        # Detener o ignorar si la fila es el resumen final (TOTAL, PRECIO, TURNOS, etc.)
                        if any(term in full_row_str for term in ["TOTAL", "PRECIO", "TURNOS", "A PAGAR", "MINUTOS CONSIDERADOS"]):
                            continue

                        # Verificar que la fila no esté completamente vacía
                        if not any(row_cells_text):
                            continue

                        # Obtener las celdas correspondientes a las columnas FIRMA
                        row_cells = table.rows[r_idx].cells
                        for col_idx in col_firma_indices:
                            if col_idx < len(row_cells):
                                cell_rect = row_cells[col_idx]
                                if cell_rect:
                                    x0, y0, x1, y1 = cell_rect
                                    w_box = x1 - x0
                                    h_box = y1 - y0
                                    y_pdf_bottom = pdf_h - y1
                                    casillas_pdf.append((x0, y_pdf_bottom, w_box, h_box))

                if casillas_pdf:
                    doc.close()
                    return casillas_pdf

            # ── ENFOQUE 2: Visión por Computadora con OpenCV (Fallback Delimitado) ──
            if cv2 and np:
                dpi = 150
                pix = page.get_pixmap(dpi=dpi)
                img_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
                if pix.n == 4:
                    img_bgr = cv2.cvtColor(img_np, cv2.COLOR_RGBA2BGR)
                else:
                    img_bgr = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

                img_h, img_w = img_bgr.shape[0], img_bgr.shape[1]
                scale_x = pdf_w / img_w
                scale_y = pdf_h / img_h

                # Ubicar límites Y estrictos de la tabla analizando palabras "FIRMA" y "TOTAL"
                text_words = page.get_text("words")
                firma_x_centers = []
                y_header_limit = 0.0
                y_footer_limit = pdf_h

                for w_info in text_words:
                    w_txt = w_info[4].strip().upper()
                    if "FIRMA" in w_txt:
                        firma_x_centers.append((w_info[0] + w_info[2]) / 2.0)
                        y_header_limit = max(y_header_limit, w_info[3])
                    elif any(t in w_txt for t in ["TOTAL", "PRECIO", "PAGAR"]):
                        if w_info[1] > y_header_limit + 20:
                            y_footer_limit = min(y_footer_limit, w_info[1])

                gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
                _, thresh = cv2.threshold(gray, 210, 255, cv2.THRESH_BINARY_INV)

                kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (max(10, img_w // 25), 1))
                detect_h = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_h)
                kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(10, img_h // 25)))
                detect_v = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_v)

                table_grid = cv2.addWeighted(detect_h, 0.5, detect_v, 0.5, 0)
                _, table_grid = cv2.threshold(table_grid, 10, 255, cv2.THRESH_BINARY)
                contours, _ = cv2.findContours(table_grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)

                for cnt in contours:
                    cx, cy, cw, ch = cv2.boundingRect(cnt)
                    x0_pdf = cx * scale_x
                    y0_pdf = cy * scale_y
                    w_pdf = cw * scale_x
                    h_pdf = ch * scale_y

                    if y0_pdf >= y_header_limit - 5 and (y0_pdf + h_pdf) <= y_footer_limit + 5:
                        center_x_pdf = x0_pdf + (w_pdf / 2.0)
                        for fx in firma_x_centers:
                            if abs(center_x_pdf - fx) <= (w_pdf / 2.0):
                                y_pdf_bottom = pdf_h - (y0_pdf + h_pdf)
                                casillas_pdf.append((x0_pdf, y_pdf_bottom, w_pdf, h_pdf))
                                break

            doc.close()
            return casillas_pdf

        except Exception as e:
            print(f"Error al detectar casillas de firma: {e}")
            return []

    @staticmethod
    def _convertir_docx_a_pdf(ruta_docx: str, ruta_pdf_salida: str) -> bool:
        """Intenta convertir un archivo DOCX a PDF usando win32com (MS Word) o docx2pdf."""
        # 1. Intentar win32com (MS Word oficial en Windows)
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc_obj = word.Documents.Open(os.path.abspath(ruta_docx))
            doc_obj.SaveAs(os.path.abspath(ruta_pdf_salida), FileFormat=17) # 17 = wdFormatPDF
            doc_obj.Close()
            word.Quit()
            if os.path.exists(ruta_pdf_salida):
                return True
        except Exception:
            pass

        # 2. Intentar librería docx2pdf
        try:
            from docx2pdf import convert
            convert(ruta_docx, ruta_pdf_salida)
            if os.path.exists(ruta_pdf_salida):
                return True
        except Exception:
            pass

        return False

    @classmethod
    def firmar_asistencia(cls, ruta_doc: str, ruta_firma: str, ruta_pdf_final: str) -> str:
        """
        Recibe un archivo .docx o .pdf de asistencia y una firma.
        Inserta la firma sin alterar el tamaño de celdas y retorna la ruta del PDF firmado final.
        """
        if not os.path.exists(ruta_doc):
            raise FileNotFoundError(f"No se encontró el archivo de asistencia: {ruta_doc}")
        if not os.path.exists(ruta_firma):
            raise FileNotFoundError(f"No se encontró la imagen de la firma: {ruta_firma}")

        ext = os.path.splitext(ruta_doc)[1].lower()

        if ext == ".docx":
            if not docx:
                raise ImportError("La librería 'python-docx' no está instalada.")

            doc = docx.Document(ruta_doc)
            firmas_insertadas = 0

            for table in doc.tables:
                col_indices = []
                if len(table.rows) > 0:
                    for idx, cell in enumerate(table.rows[0].cells):
                        if "FIRMA" in cell.text.upper():
                            col_indices.append(idx)

                for row in table.rows[1:]:
                    for col_idx in col_indices:
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            cell.text = "" # Limpiar texto manteniendo formato
                            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            run = p.add_run()
                            # Firma de tamaño exacto que encaja en la celda
                            run.add_picture(ruta_firma, width=Inches(0.95), height=Inches(0.28))
                            firmas_insertadas += 1

            if firmas_insertadas == 0:
                for p in doc.paragraphs:
                    if "FIRMA" in p.text.upper():
                        p.text = ""
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run()
                        run.add_picture(ruta_firma, width=Inches(1.2), height=Inches(0.35))

            # Guardar temporalmente el docx firmado
            temp_docx = ruta_pdf_final.replace(".pdf", "_temp.docx")
            doc.save(temp_docx)

            # Convertir a PDF (Salida obligatoria en PDF)
            exito_pdf = cls._convertir_docx_a_pdf(temp_docx, ruta_pdf_final)

            # Limpiar temporal
            if os.path.exists(temp_docx):
                try: os.remove(temp_docx)
                except Exception: pass

            if exito_pdf:
                return ruta_pdf_final
            else:
                raise RuntimeError("No se pudo convertir el DOCX a PDF. Instala MS Word o 'docx2pdf'.")

        elif ext == ".pdf":
            if not pypdf:
                raise ImportError("La librería 'pypdf' no está instalada.")

            reader = pypdf.PdfReader(ruta_doc)
            writer = pypdf.PdfWriter()

            packet = io.BytesIO()
            can = canvas.Canvas(packet)

            page0 = reader.pages[0]
            w = float(page0.mediabox.width)
            h = float(page0.mediabox.height)

            # Detección dinámica de casillas de firma por Visión por Computadora / PyMuPDF Table
            casillas_cv = cls._detectar_casillas_firma_cv(ruta_doc)

            if casillas_cv:
                # Estampar la firma perfectamente centrada en cada celda detectada dentro de la tabla
                for (x_pdf, y_pdf_bottom, w_pdf, h_pdf) in casillas_cv:
                    box_w = w_pdf * 0.82
                    box_h = h_pdf * 0.75

                    center_x = x_pdf + (w_pdf / 2.0)
                    center_y = y_pdf_bottom + (h_pdf / 2.0)

                    pos_x = center_x - (box_w / 2.0)
                    pos_y = center_y - (box_h / 2.0)

                    can.drawImage(
                        ruta_firma, pos_x, pos_y,
                        width=box_w, height=box_h,
                        mask='auto', preserveAspectRatio=True
                    )
            else:
                # Fallback proporcional si no se detectan celdas por tabla
                y_starts = [h * 0.82 - (i * 26.5) for i in range(12)]
                for y_pos in y_starts:
                    if y_pos > 150:
                        can.drawImage(ruta_firma, w * 0.31, y_pos + 2, width=58, height=18, mask='auto', preserveAspectRatio=True)
                        can.drawImage(ruta_firma, w * 0.57, y_pos + 2, width=58, height=18, mask='auto', preserveAspectRatio=True)

            can.save()
            packet.seek(0)

            overlay_pdf = pypdf.PdfReader(packet)
            overlay_page = overlay_pdf.pages[0]

            for i, page in enumerate(reader.pages):
                if i == 0:
                    page.merge_page(overlay_page)
                writer.add_page(page)

            with open(ruta_pdf_final, "wb") as f_out:
                writer.write(f_out)

            return ruta_pdf_final
        else:
            raise ValueError("Formato de entrada no soportado. Debe ser .docx o .pdf")
